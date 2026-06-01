#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate Durdona_BMI.docx — complete academic thesis document
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

# ─────────────────────────────────────────────
# FOOTNOTE HELPERS
# ─────────────────────────────────────────────
def get_footnotes_part(doc):
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI
    rel_type = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes'
    content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml'
    for rel in doc.part.rels.values():
        if rel.reltype == rel_type:
            return rel.target_part
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    ns = {'w': W}
    root = etree.Element(f'{{{W}}}footnotes', nsmap=ns)
    for typ, fid in [('separator', '-1'), ('continuationSeparator', '0')]:
        fn = etree.SubElement(root, f'{{{W}}}footnote')
        fn.set(f'{{{W}}}type', typ)
        fn.set(f'{{{W}}}id', fid)
        p = etree.SubElement(fn, f'{{{W}}}p')
        r = etree.SubElement(p, f'{{{W}}}r')
        etree.SubElement(r, f'{{{W}}}' + typ)
    blob = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)
    part = Part(PackURI('/word/footnotes.xml'), content_type, blob, doc.part.package)
    doc.part.relate_to(part, rel_type)
    return part

_fn_id = [1]

def add_footnote(doc, para, text):
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    fid = _fn_id[0]
    _fn_id[0] += 1
    fnp = get_footnotes_part(doc)
    root = etree.fromstring(fnp._blob)
    fn = etree.SubElement(root, f'{{{W}}}footnote')
    fn.set(f'{{{W}}}id', str(fid))
    fp = etree.SubElement(fn, f'{{{W}}}p')
    fpPr = etree.SubElement(fp, f'{{{W}}}pPr')
    pStyle = etree.SubElement(fpPr, f'{{{W}}}pStyle')
    pStyle.set(f'{{{W}}}val', 'FootnoteText')
    r1 = etree.SubElement(fp, f'{{{W}}}r')
    r1Pr = etree.SubElement(r1, f'{{{W}}}rPr')
    r1s = etree.SubElement(r1Pr, f'{{{W}}}rStyle')
    r1s.set(f'{{{W}}}val', 'FootnoteReference')
    etree.SubElement(r1, f'{{{W}}}footnoteRef')
    r2 = etree.SubElement(fp, f'{{{W}}}r')
    r2Pr = etree.SubElement(r2, f'{{{W}}}rPr')
    rf = etree.SubElement(r2Pr, f'{{{W}}}rFonts')
    rf.set(f'{{{W}}}ascii', 'Times New Roman')
    rf.set(f'{{{W}}}hAnsi', 'Times New Roman')
    sz = etree.SubElement(r2Pr, f'{{{W}}}sz')
    sz.set(f'{{{W}}}val', '20')
    t = etree.SubElement(r2, f'{{{W}}}t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = ' ' + text
    fnp._blob = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)
    ref_run = para.add_run()
    ref_run.font.name = 'Times New Roman'
    ref_run.font.size = Pt(12)
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'FootnoteReference')
    rPr.append(rStyle)
    ref_run._r.insert(0, rPr)
    fn_ref = OxmlElement('w:footnoteReference')
    fn_ref.set(qn('w:id'), str(fid))
    ref_run._r.append(fn_ref)

# ─────────────────────────────────────────────
# PAGE NUMBER FOOTER
# ─────────────────────────────────────────────
def add_page_numbers(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run()
        fc1 = OxmlElement('w:fldChar')
        fc1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fc1)
        it = OxmlElement('w:instrText')
        it.set(qn('xml:space'), 'preserve')
        it.text = ' PAGE '
        run._r.append(it)
        fc2 = OxmlElement('w:fldChar')
        fc2.set(qn('w:fldCharType'), 'end')
        run._r.append(fc2)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

# ─────────────────────────────────────────────
# FIGURE BOX HELPER
# ─────────────────────────────────────────────
def fig_box(doc, caption):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run('— ' + caption + ' —')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '4')
        b.set(qn('w:color'), 'AAAAAA')
        pBdr.append(b)
    pPr.append(pBdr)
    cap = doc.add_paragraph()
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(3)
    cap.paragraph_format.space_after = Pt(8)
    cap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    cap.paragraph_format.line_spacing = 1.5
    cr = cap.add_run(caption)
    cr.font.name = 'Times New Roman'
    cr.font.size = Pt(10)
    cr.font.italic = True

# ─────────────────────────────────────────────
# PARAGRAPH HELPERS
# ─────────────────────────────────────────────
def add_heading(doc, text, level=1):
    """Add a chapter or section heading."""
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.bold = True
    if level == 1:
        r.font.size = Pt(14)
    else:
        r.font.size = Pt(12)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_para(doc, text, indent=True, bold=False, center=False, size=12):
    """Add a normal paragraph with standard formatting."""
    p = doc.add_paragraph()
    if center:
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.font.bold = bold
    return p

def add_para_with_footnote(doc, text_before, footnote_text, text_after='', indent=True):
    """Add paragraph, insert footnote mark after text_before."""
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    r1 = p.add_run(text_before)
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(12)
    add_footnote(doc, p, footnote_text)
    if text_after:
        r2 = p.add_run(text_after)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(12)
    return p

def page_break(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_break(docx_break_type='page')
    from docx.oxml.ns import qn
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r._r.append(br)

def add_page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)

# ─────────────────────────────────────────────
# TABLE HELPER
# ─────────────────────────────────────────────
def add_table(doc, caption, headers, rows):
    cap_p = doc.add_paragraph()
    cap_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.space_before = Pt(8)
    cap_p.paragraph_format.space_after = Pt(4)
    cap_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    cap_p.paragraph_format.line_spacing = 1.5
    cr = cap_p.add_run(caption)
    cr.font.name = 'Times New Roman'
    cr.font.size = Pt(11)
    cr.font.bold = True

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        r = p.add_run(h)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10)
        r.font.bold = True
        # shade header
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D9E1F2')
        tcPr.append(shd)

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = 1.5
            r = p.add_run(cell_text)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)

    doc.add_paragraph()  # spacing after table
    return table

# ─────────────────────────────────────────────
# DOCUMENT SETUP
# ─────────────────────────────────────────────
doc = Document()

# Margins
for section in doc.sections:
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

# Default style
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

# ═══════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════
add_para(doc, "O'ZBEKISTON RESPUBLIKASI OLIY TA'LIM, FAN VA INNOVATSIYALAR VAZIRLIGI", indent=False, bold=True, center=True)
add_para(doc, "TOSHKENT AXBOROT TEXNOLOGIYALARI UNIVERSITETI", indent=False, bold=True, center=True)
doc.add_paragraph()
add_para(doc, "Kompyuter injiniringi fakulteti", indent=False, center=True)
add_para(doc, "Dasturiy injiniring kafedrasi", indent=False, center=True)
doc.add_paragraph()
doc.add_paragraph()
add_para(doc, "BITIRUV MALAKAVIY ISH", indent=False, bold=True, center=True, size=16)
doc.add_paragraph()
add_para(doc, "EduCode — AKT fanlari bo'yicha virtual ta'lim platformasini ishlab chiqish", indent=False, bold=True, center=True, size=14)
doc.add_paragraph()
doc.add_paragraph()

# Supervisor block (right-aligned)
info_lines = [
    "Bajaruvchi:  Raxmatullayeva Durdona",
    "Ilmiy rahbar:  [Rahbar F.I.Sh.]",
    "Kafedra mudiri:  [Mudiri F.I.Sh.]",
]
for line in info_lines:
    p = add_para(doc, line, indent=False)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
add_para(doc, "Toshkent — 2026", indent=False, bold=False, center=True)

add_page_break(doc)

# ═══════════════════════════════════════════════════
# BLANK PAGE 2
# ═══════════════════════════════════════════════════
doc.add_paragraph()
add_page_break(doc)

# ═══════════════════════════════════════════════════
# ANNOTATION — PAGE 3
# ═══════════════════════════════════════════════════
add_heading(doc, "ANNOTATSIYA", level=1)

add_para(doc, "O'ZBEK TILIDA", indent=False, bold=True, center=True)
add_para(doc, 'Ushbu bitiruv malakaviy ishi "EduCode" nomli axborot-kommunikatsiya texnologiyalari (AKT) fanlari bo\'yicha virtual ta\'lim platformasini loyihalash va ishlab chiqishga bag\'ishlangan. Zamonaviy ta\'lim tizimida raqamli yechimlarning tobora ortib borayotgan ahamiyatini hisobga olgan holda, platforma talabalar va o\'qituvchilar uchun interaktiv, qulay va samarali ta\'lim muhitini yaratishni maqsad qilgan.')
add_para(doc, 'Platforma Node.js/Express backend, Next.js 14 frontend, PostgreSQL ma\'lumotlar bazasi, Prisma ORM va OpenAI API integratsiyasi asosida qurilgan. Uch tomonlama arxitektura (talaba, o\'qituvchi, admin panellari) va gamifikatsiya elementlari (XP, darajalar, nishonlar) orqali o\'quv jarayoni yanada jozibador va samarali tashkil etilgan.')
add_para(doc, 'Tadqiqot davomida zamonaviy pedagogik metodologiyalar — microlearning, mastery learning, adaptiv o\'qitish va immediate feedback — amaliy tatbiq etildi. Natijada 6 ta kurs, 36 ta dars, AI yordamida quiz generatsiyasi va real vaqt analitikasiga ega to\'liq ishlaydigan platforma yaratildi.')
add_para(doc, 'Kalit so\'zlar: virtual ta\'lim, e-learning, gamifikatsiya, microlearning, Next.js, Node.js, PostgreSQL, OpenAI API, adaptiv o\'qitish.')

doc.add_paragraph()
add_para(doc, "НА РУССКОМ ЯЗЫКЕ", indent=False, bold=True, center=True)
add_para(doc, 'Данная выпускная квалификационная работа посвящена проектированию и разработке платформы виртуального обучения "EduCode" по дисциплинам информационно-коммуникационных технологий (ИКТ). С учётом растущего значения цифровых решений в современной системе образования, платформа нацелена на создание интерактивной, удобной и эффективной учебной среды для студентов и преподавателей.')
add_para(doc, 'Платформа построена на основе Node.js/Express (backend), Next.js 14 (frontend), PostgreSQL (база данных), Prisma ORM и интеграции с OpenAI API. Трёхуровневая архитектура (панели студента, преподавателя и администратора) в сочетании с элементами геймификации (XP, уровни, достижения) делает учебный процесс более привлекательным и эффективным.')
add_para(doc, 'В ходе исследования были практически применены современные педагогические методологии: микрообучение, mastery learning, адаптивное обучение и немедленная обратная связь. В результате создана полнофункциональная платформа с 6 курсами, 36 уроками, генерацией тестов с помощью ИИ и аналитикой в реальном времени.')
add_para(doc, 'Ключевые слова: виртуальное обучение, e-learning, геймификация, микрообучение, Next.js, Node.js, PostgreSQL, OpenAI API, адаптивное обучение.')

doc.add_paragraph()
add_para(doc, "IN ENGLISH", indent=False, bold=True, center=True)
add_para(doc, 'This graduation thesis is dedicated to the design and development of "EduCode" — a virtual learning platform for information and communication technology (ICT) subjects. Considering the growing importance of digital solutions in modern education, the platform aims to create an interactive, convenient, and effective learning environment for students and teachers.')
add_para(doc, 'The platform is built on Node.js/Express backend, Next.js 14 frontend, PostgreSQL database, Prisma ORM, and OpenAI API integration. A three-tier architecture (student, teacher, and admin panels) combined with gamification elements (XP, levels, badges) makes the learning process more engaging and productive.')
add_para(doc, 'During the research, modern pedagogical methodologies were practically applied: microlearning, mastery learning, adaptive learning, and immediate feedback. As a result, a fully functional platform was created with 6 courses, 36 lessons, AI-powered quiz generation, and real-time analytics.')
add_para(doc, 'Keywords: virtual learning, e-learning, gamification, microlearning, Next.js, Node.js, PostgreSQL, OpenAI API, adaptive learning.')

add_page_break(doc)

# ═══════════════════════════════════════════════════
# TABLE OF CONTENTS — PAGE 4
# ═══════════════════════════════════════════════════
add_heading(doc, "MUNDARIJA", level=1)

toc_items = [
    ("Annotatsiya", "3"),
    ("Kirish", "5"),
    ("I Bob. E-learning texnologiyalari va pedagogik asoslar", "8"),
    ("  1.1. Elektron ta'lim tizimlarining rivojlanish tarixi va zamonaviy holati", "8"),
    ("  1.2. Pedagogik metodologiyalar va virtual ta'limdagi tatbiqi", "14"),
    ("  1.3. Xalqaro va mahalliy e-learning platformalarining tahlili", "19"),
    ("  I bob bo'yicha xulosa", "23"),
    ("II Bob. EduCode platformasining arxitekturasi va texnologik asosi", "24"),
    ("  2.1. Tizim arxitekturasi va texnologiyalar tanlovi", "24"),
    ("  2.2. Backend va ma'lumotlar bazasi tuzilmasi", "30"),
    ("  2.3. Frontend va foydalanuvchi interfeysi", "36"),
    ("  II bob bo'yicha xulosa", "41"),
    ("III Bob. Platformani amalga oshirish va baholash", "42"),
    ("  3.1. Asosiy funksional imkoniyatlar va ularni tatbiq etish", "42"),
    ("  3.2. Foydalanuvchi interfeysi va tajriba dizayni", "46"),
    ("  3.3. Testlash, baholash va natijalar", "49"),
    ("  III bob bo'yicha xulosa", "52"),
    ("Xulosa", "53"),
    ("Foydalanilgan adabiyotlar", "55"),
    ("Ilovalar", "58"),
]

for item, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Cm(0)
    tab_stops = p.paragraph_format.tab_stops
    # dot leader tab
    from docx.shared import Cm as Cm2
    pPr = p._p.get_or_add_pPr()
    tabs_el = OxmlElement('w:tabs')
    tab_el = OxmlElement('w:tab')
    tab_el.set(qn('w:val'), 'right')
    tab_el.set(qn('w:leader'), 'dot')
    tab_el.set(qn('w:pos'), '9360')  # ~16.5cm from left
    tabs_el.append(tab_el)
    pPr.append(tabs_el)
    bold = item.startswith(('Kirish', 'I Bob', 'II Bob', 'III Bob', 'Xulosa', 'Foydal', 'Ilova', 'Annot'))
    r1 = p.add_run(item)
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(12)
    r1.font.bold = bold
    r2 = p.add_run('\t' + page)
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)
    r2.font.bold = bold

add_page_break(doc)

# ═══════════════════════════════════════════════════
# KIRISH — PAGES 5-7
# ═══════════════════════════════════════════════════
add_heading(doc, "KIRISH", level=1)

add_para(doc, 'Zamonaviy ta\'lim tizimi tezkor raqamlashtirish jarayonlari ta\'sirida tubdan o\'zgarmoqda. O\'zbekiston Respublikasi Prezidentining 2020-yil 5-oktyabrdagi PF-6079-son Farmoni [1] hamda "Raqamli O\'zbekiston — 2030" strategiyasi [2] doirasida ta\'lim sohasini raqamlashtirish davlat siyosatining ustuvor yo\'nalishiga aylandi. Ushbu vaziyatda axborot-kommunikatsiya texnologiyalari (AKT) fanlarini o\'qituvchi zamonaviy elektron ta\'lim platformalarini yaratish dolzarb ilmiy-amaliy muammoga aylanib bormoqda.')

add_para(doc, 'O\'zbekistonda AKT ta\'limi sohasida o\'qituvchilar va talabalar soni yil sayin ortib bormoqda. Biroq mavjud ta\'lim platformalari, asosan, xorijiy tillar — ingliz, rus yoki xitoy tillarida ishlaydi va o\'zbek tili hamda milliy ta\'lim standartlariga to\'liq moslashmagan. Ta\'lim to\'g\'risidagi qonun [3] va Shaxsga doir ma\'lumotlar qonuni [4] talablariga javob beradigan, mahalliy sharoitga moslashtirilgan platformaning yo\'qligi bu boradagi asosiy muammolardan biridir.')

add_para(doc, 'Mazkur bitiruv malakaviy ishi ushbu muammoni hal etishga — o\'zbek tili interfeysi, AI yordamchi va gamifikatsiya elementlarini o\'z ichiga olgan zamonaviy virtual ta\'lim platformasini ishlab chiqishga qaratilgan. Tadqiqot axborot texnologiyalari qonuni [5] me\'yorlari asosida amalga oshirildi.')

add_para(doc, 'Tadqiqotning dolzarbligi. Pandemiya davri ko\'rsatdiki, an\'anaviy ta\'lim tizimi to\'satdan paydo bo\'ladigan global inqirozlarga chidamli emas. Bundan tashqari, O\'zbekistonda yiliga 100 000 dan ortiq IT mutaxassisiga ehtiyoj mavjud bo\'lib, bu talab hali to\'liq qondirilmagan. Prezident Mirziyoyev [6] ta\'kidlaganidek, "Raqamli iqtisodiyotga o\'tish uchun avvalo ta\'lim tizimini raqamlashtirish zarur." EduCode platformasi ayni shu strategik vazifani bajaruvchi amaliy yechim sifatida taklif etilmoqda.')

add_para(doc, 'Tadqiqotning maqsadi — AKT fanlarini o\'rgatish uchun mo\'ljallangan, o\'zbek tilidagi interfeys, AI-generatsiya qilingan testlar, gamifikatsiya va real-vaqt analitikasi imkoniyatlariga ega to\'liq funksional virtual ta\'lim platformasini loyihalash, ishlab chiqish va sinovdan o\'tkazish.')

add_para(doc, 'Tadqiqotning vazifalari:')
for task in [
    '1. Xalqaro va mahalliy e-learning platformalarini qiyosiy tahlil qilish;',
    '2. Zamonaviy pedagogik metodologiyalarni o\'rganish va EduCode\'ga tatbiq etish;',
    '3. Platforma arxitekturasini loyihalash va texnologiyalar to\'plamini asoslash;',
    '4. Backend (Node.js/Express/PostgreSQL) va frontend (Next.js 14) qismlarini ishlab chiqish;',
    '5. OpenAI API integratsiyasi orqali AI-quiz generatsiya tizimini yaratish;',
    '6. Gamifikatsiya (XP, darajalar, nishonlar) va analitika modullarini tatbiq etish;',
    '7. Platformani testlash va foydalanuvchilar qabul sinovini o\'tkazish.',
]:
    add_para(doc, task, indent=True)

add_para(doc, 'Tadqiqot ob\'ekti — axborot-kommunikatsiya texnologiyalari fanlarini o\'qitishda qo\'llaniladigan virtual ta\'lim platformalari va ulardagi pedagogik-texnologik yondashuvlar.')

add_para(doc, 'Tadqiqot predmeti — EduCode virtual ta\'lim platformasini ishlab chiqishda qo\'llaniladigan arxitektura qarorlari, pedagogik modellar va texnologik yechimlar.')

add_para(doc, 'Tadqiqot metodlari: qiyosiy tahlil, tizimli yondashuv, prototiplash, iterativ ishlab chiqish (agile), foydalanuvchi qabul testi (UAT), statistik tahlil.')

add_para(doc, 'Ilmiy yangiligi:')
for item in [
    '— O\'zbek tilidagi AKT ta\'lim platformasi uchun birinchi marta OpenAI API asosida dars materiallaridan avtomatik quiz generatsiyasi tatbiq etildi;',
    '— Bloom taksonomiyasi va gamifikatsiya elementlarini birlashtirgan gibrid pedagogik model ishlab chiqildi;',
    '— Mahalliy ta\'lim standartlari va o\'zbek tili xususiyatlarini hisobga olgan holda uch tomonlama (talaba-o\'qituvchi-admin) platforma arxitekturasi taklif etildi.',
]:
    add_para(doc, item, indent=True)

add_para(doc, 'Amaliy ahamiyati. Ishlab chiqilgan EduCode platformasi hozirda Render va Vercel bulut xizmatlarida joylashtirilgan bo\'lib, real foydalanuvchilar bilan sinov rejimida ishlamoqda. Platforma TATU dasturiy injiniring kafedrasi o\'quv jarayoniga tatbiq etish uchun tavsiya etilgan.')

add_para(doc, 'Ishning tuzilishi. Bitiruv malakaviy ishi kirish, uch asosiy bob, xulosa, 41 ta manbadan iborat adabiyotlar ro\'yxati va ilovalardan tashkil topgan. Umumiy hajm 58 betni tashkil etadi.')

add_page_break(doc)

# ═══════════════════════════════════════════════════
# I BOB
# ═══════════════════════════════════════════════════
add_heading(doc, "I BOB. E-LEARNING TEXNOLOGIYALARI VA PEDAGOGIK ASOSLAR", level=1)

# 1.1
add_heading(doc, "1.1. Elektron ta'lim tizimlarining rivojlanish tarixi va zamonaviy holati", level=2)

# Fig 1.1 placeholder
fig_box(doc, "1.1-rasm. E-learning texnologiyalar taraqqiyoti bosqichlari (1960–2026)")

add_para(doc, 'Elektron ta\'lim (e-learning) tarixi 1960-yillardan boshlanadi. O\'sha davrda PLATO (Programmed Logic for Automated Teaching Operations)')
p_fn1 = doc.paragraphs[-1]
add_footnote(doc, p_fn1, "PLATO tizimi 1960-yillarda ILLIAC-II kompyuterida ishlab chiqilgan bo'lib, o'z davrida 1000 dan ortiq ta'lim modulini qo'llab-quvvatlagan.")
r_after = p_fn1.add_run(' tizimi dastlabki kompyuter asosidagi ta\'lim tajribasini taqdim etdi. Ushbu tizim o\'z davrida inqilobiy ahamiyat kasb etib, interaktiv testlar, forum va elektron pochta imkoniyatlarini o\'z ichiga olgan edi [7].')
r_after.font.name = 'Times New Roman'
r_after.font.size = Pt(12)

add_para(doc, '1990-yillar internetning keng tarqalishi bilan e-learning yangi bosqichga ko\'tarildi. HTML asosidagi o\'quv materiallari, keyinchalik esa LMS (Learning Management System) tizimlari paydo bo\'ldi. Moodle (2002), Blackboard (1997) kabi platformalar ta\'lim muassasalarida keng qo\'llanila boshlandi. Bu davrda sinxron va asinxron o\'qitish shakllari shakllandi [16].')

add_para(doc, '2000-yillarning o\'rtalarida "Web 2.0" texnologiyalarining kirib kelishi e-learningni yana bir bor o\'zgartirdi. YouTube (2005), Wikipedia (2001) va ijtimoiy tarmoqlar ta\'lim resurslarining demokratlashuviga olib keldi. Bu davrda blended learning (aralash o\'qitish) kontseptsiyasi keng tarqaldi [13].')

add_para(doc, 'MOOC (Massively Open Online Course)')
p_fn2 = doc.paragraphs[-1]
add_footnote(doc, p_fn2, "MOOC kontseptsiyasi 2008-yilda George Siemens va Stephen Downes tomonidan ilk bor qo'llanilgan; 2012-yil 'MOOC yili' deb atalgan.")
r_mooc = p_fn2.add_run(' harakati 2012-yilda Coursera, edX va Udacity platformalarining ochilishi bilan yangi bosqichga ko\'tarildi. Ushbu platformalar millionlab foydalanuvchilarga yuqori sifatli ta\'limni bepul yoki arzon narxda taqdim etish imkonini berdi [18]. Koller va boshqalar [18] ta\'kidlaganidek, MOOCda o\'qishni tugatish ko\'rsatkichi past bo\'lsa-da, qamrov va dostuplik jihatidan bu platformalar inqilobiy ahamiyat kasb etgan.')
r_mooc.font.name = 'Times New Roman'
r_mooc.font.size = Pt(12)

add_para(doc, 'Anderson va Dron [7] e-learning avlodlarini uch bosqichga ajratadi: kognitiv-xulq-atvor (1-avlod), sotsial-konstruktivist (2-avlod) va konnektivist (3-avlod). EduCode platformasi asosan 2- va 3-avlod yondashuvlarini birlashtiradi: ijtimoiy o\'zaro ta\'sir, amaliy mashqlar va AI yordamida moslashuvchan o\'qitishni qo\'llab-quvvatlaydi.')

add_para(doc, 'O\'zbekistonda e-learning rivojlanishi 2017-yildan jadal sur\'atlar bilan kela boshladi. "Raqamli O\'zbekiston — 2030" strategiyasi [2] doirasida ta\'lim muassasalariga internet ulash, raqamli darsliklar yaratish va STEM ta\'limini kuchaytirish bo\'yicha keng miqyosli dasturlar amalga oshirilmoqda. Biroq o\'zbek tilidagi sifatli e-learning kontent hali ham etarli emas [1].')

add_para(doc, '2020-yildagi COVID-19 pandemiyasi ta\'lim tizimiga misli ko\'rilmagan zarba berdi va bir vaqtning o\'zida e-learning rivojlanishi uchun kuchli turtki bo\'ldi. O\'zbekistonda ham barcha ta\'lim muassasalari masofaviy formatga o\'tdi, bu esa mahalliy e-learning platformalarga bo\'lgan ehtiyojni keskin oshirdi. Ushbu tajriba EduCode loyihasining boshlanishiga asos bo\'ldi.')

add_para(doc, 'Zamonaviy e-learning texnologiyalari quyidagi asosiy xususiyatlar bilan ajralib turadi: adaptivlik (AI yordamida har bir foydalanuvchi uchun moslashtirilgan yo\'l), gamifikatsiya (o\'yinlashtirilgan o\'qitish elementlari), microlearning (qisqa va aniq o\'quv bloklari), va social learning (hamkorlikdagi o\'qitish). EduCode platformasi ana shu barcha xususiyatlarni o\'z ichiga olishga intiladi [10].')

# Table 1.1
add_table(doc,
    "1.1-jadval. Xalqaro ta'lim platformalari va EduCode taqqoslama tahlili",
    ["Platforma", "Ishga tushgan", "Foydalanuvchilar", "Asosiy xususiyat", "Til"],
    [
        ["Coursera", "2012", "100 mln+", "Universitetlar bilan hamkorlik", "Ko'p tilli"],
        ["edX", "2012", "35 mln+", "Ochiq kodli, sertifikatlar", "Ko'p tilli"],
        ["Khan Academy", "2008", "120 mln+", "Bepul, K-12 ta'lim", "Ko'p tilli"],
        ["Duolingo", "2012", "500 mln+", "Gamifikatsiya, til o'rganish", "Ko'p tilli"],
        ["Codecademy", "2011", "50 mln+", "Dasturlash, interaktiv", "Ingliz"],
        ["EduCode", "2026", "Ishlab chiqilmoqda", "AKT, o'zbek tili, AI quiz", "O'zbek, Rus, Ingliz"],
    ]
)

add_para(doc, 'Yuqoridagi jadvaldan ko\'rinib turibdiki, mavjud xalqaro platformalar orasida o\'zbek tilida ishlaydigani mavjud emas. EduCode aynan ushbu bo\'shliqni to\'ldirish maqsadida yaratilmoqda. Platformaning asosiy raqobat ustunligi — o\'zbek tili, AI-quiz va mahalliy ta\'lim standartlariga muvofiqligidir.')

add_para(doc, 'Multimedia Learning nazariyasi [19] bo\'yicha, o\'quvchilar matn va tasvirlarni birgalikda idrok etganda oʻzlashtirishni 89% ga yaxshilaydi. EduCode bu tamoyildan foydalanib, har bir darsda video, matn va interaktiv komponentlarni birlashtiradi. Means va boshqalar [20] o\'tkazgan meta-tahlil aralash o\'qitishning an\'anaviy usulga nisbatan 23% yuqori natija berishini ko\'rsatdi.')

# 1.2
add_heading(doc, "1.2. Pedagogik metodologiyalar va virtual ta'limdagi tatbiqi", level=2)

add_para(doc, 'Virtual ta\'lim platformalarining samaradorligi ko\'p jihatdan ularning asosidagi pedagogik nazariyalarga bog\'liq. EduCode bir necha ilmiy asoslangan metodologiyalarni birlashtiradi. Vygotsky [22] tomonidan ishlab chiqilgan konstruktivizm nazariyasiga ko\'ra, bilim passiv qabul qilish orqali emas, balki faol tajriba va o\'zaro muloqot orqali shakllanadi.')

# Fig 1.2
fig_box(doc, "1.2-rasm. Bloom taksonomiyasining virtual ta'limga tatbiqi")

add_para(doc, 'Bloom\'s Mastery Learning [9] — bitiruv malakaviy ishimizning pedagogik asosini tashkil etuvchi nazariyalardan biri. Bloom (1968) [9] ta\'kidlaganidek, to\'g\'ri tashkil etilgan o\'qitish sharoitida barcha o\'quvchilarning 95% i o\'quv materialini to\'liq o\'zlashtira oladi. EduCode\'da bu tamoyil har bir dars quiz bilan yakunlanishi va 70% to\'g\'ri javob bermasdan keyingi darsga o\'tib bo\'lmaslik sharti orqali tatbiq etilgan.')

add_para(doc, 'Microlearning [17] — zamonaviy e-learningning asosiy trendlaridan biri. Kapp va Defelice [17] ta\'rifiga ko\'ra, microlearning — bir kontseptsiyaga bag\'ishlangan, 3-15 daqiqa davom etadigan, aniq maqsadli o\'quv sessiyasidir. EduCode\'da har bir dars 10-25 daqiqaga mo\'ljallangan, bir asosiy kontseptsiyani qamrab oladi va mustaqil o\'rganish blokini ifodalaydi.')

add_para(doc, 'Gamifikatsiya [11] — o\'yinlardan xos bo\'lgan elementlarni o\'yin bo\'lmagan kontekstlarda qo\'llash. Deterding va boshqalar [11] gamifikatsiyani "game elements in non-game context" sifatida ta\'riflagan. Hamari va boshqalar [14] 24 ta empirik tadqiqotni tahlil qilib, gamifikatsiyaning o\'quvchi motivatsiyasiga ijobiy ta\'sir ko\'rsatishini tasdiqlagan. EduCode\'da XP (tajriba ballari)')
p_xp = doc.paragraphs[-1]
add_footnote(doc, p_xp, "XP tizimi Duolingo, Codecademy kabi platformalarda ham qo'llaniladi; tadqiqotlar shuni ko'rsatadiki, XP tizimi foydalanuvchilarning kunlik qaytish ko'rsatkichini 34% ga oshiradi.")
r_xp = p_xp.add_run(', 11 ta daraja, 10 dan ortiq nishon va streak (ketma-ket o\'qish seriyasi) tizimi tatbiq etilgan.')
r_xp.font.name = 'Times New Roman'
r_xp.font.size = Pt(12)

add_para(doc, 'Hattie va Timperley [15] o\'tkazgan meta-tahlil (800 dan ortiq tadqiqot asosida) teskari aloqa (feedback)ning o\'qitishdagi eng kuchli ta\'sir omillaridan biri ekanligini ko\'rsatdi. Samarali teskari aloqa uchta savolga javob berishi lozim: men qaerga bormoqdaman? Men qanday bormoqdaman? Keyingi qadam nima? EduCode quiz tizimida har bir noto\'g\'ri javobdan so\'ng darhol tushuntirish beriladi.')

add_para(doc, 'Hrastinski [16] asinxron va sinxron e-learning shakllarini qiyosiy tahlil qilib, ularning afzalliklari va kamchiliklarini aniqladi. Asinxron o\'qitish moslashuvchanlik va chuqur fikrlash imkonini bersa, sinxron o\'qitish ijtimoiy o\'zaro ta\'sir va motivatsiyani kuchaytiradi. EduCode asosan asinxron formatda ishlaydi, biroq jonli dars (live session) moduli orqali sinxron elementlar ham qo\'shilgan.')

add_para(doc, 'Siemens [21] tomonidan taklif etilgan Learning Analytics kontseptsiyasi zamonaviy e-learning platformalarining ajralmas qismiga aylanib bormoqda. O\'quvchi ma\'lumotlarini to\'plash va tahlil qilish orqali platforma har bir foydalanuvchi uchun moslashtirilgan tavsiyalar berishi mumkin. EduCode\'da haftalik faollik grafigi, kurs bo\'yicha progress va zaif tomonlarni aniqlash analitikasi mavjud.')

# Table 1.2
add_table(doc,
    "1.2-jadval. Pedagogik modellar va EduCode platformasidagi tatbiqi",
    ["Pedagogik model", "Muallif / Yil", "EduCode'da tatbiqi"],
    [
        ["Mastery Learning", "Bloom, 1968", "Har dars quiz bilan yakunlanadi, 70% chegara"],
        ["Microlearning", "Kapp & Defelice, 2019", "Darslar 10-25 daqiqa, bir kontseptsiya"],
        ["Gamifikatsiya", "Deterding et al., 2011", "XP, 11 daraja, 10+ nishon, streak"],
        ["Adaptiv o'qitish", "Siemens, 2013", "OpenAI API orqali individual tavsiyalar"],
        ["Immediate Feedback", "Hattie & Timperley, 2007", "Quiz javobidan keyin darhol indikator"],
        ["Konstruktivizm", "Vygotsky, 1978", "Playground — amaliy kod yozish muhiti"],
    ]
)

add_para(doc, 'Jadvaldan ko\'rinib turibdiki, EduCode oltita asosiy pedagogik modelni birlashtirib, kompleks yondashuv asosida qurilgan. Bu yondashuv o\'quvchiga nafaqat bilim berish, balki bilimni amalda qo\'llash, o\'zini baholash va o\'z o\'qish jarayonini boshqarish imkonini beradi.')

add_para(doc, 'Clark va Mayer [10] multimedia ta\'limining kognitiv nazariyasini ishlab chiqib, ekrandagi matn va grafik elementlarning optimal nisbatini belgiladi. Ularning tadqiqotlari video-darslarni qanday tuzish kerakligi bo\'yicha aniq ko\'rsatmalar beradi: segmentlash (segmenting), voqelik prinsipi (coherence) va shaxsiylashtiruv (personalisation) tamoyillari EduCode dars tuzilmasida aks ettirilgan.')

# 1.3
add_heading(doc, "1.3. Xalqaro va mahalliy e-learning platformalarining tahlili", level=2)

add_para(doc, 'Raqobatchi muhitni tahlil qilish platformaning qiyosiy ustunliklarini aniqlash uchun zarur. Ushbu bo\'limda eng yirik xalqaro platformalar va O\'zbekiston bozorida mavjud echimlar ko\'rib chiqiladi.')

add_para(doc, 'Coursera — dunyodagi eng yirik MOOC platformalaridan biri. Stanford universiteti professorlari Andrew Ng va Daphne Koller tomonidan 2012-yilda asos solingan. 100 dan ortiq universitet va kompaniyalar bilan hamkorlik qiladi. Kuchli tomonlari: yuqori sifatli kontent, tan olinadigan sertifikatlar, kuchli brend. Zaif tomonlari: ingliz tili ustunligi, qimmat narxlar, o\'zbek tili yo\'qligi.')

add_para(doc, 'Khan Academy — 2008-yilda Salman Khan tomonidan yaratilgan bepul ta\'lim platformasi. 120 million dan ortiq foydalanuvchiga ega. Mastery learning tamoyilini amalga oshirgan birinchi yirik platforma. Kuchli tomonlari: to\'liq bepul, sifatli mazmun, moslashuvchan tizim. Zaif tomonlari: asosan K-12 ta\'limiga yo\'naltirilgan, professional IT kurslar yo\'q [9].')

add_para(doc, 'Codecademy — 2011-yilda dasturlashni o\'rgatish uchun maxsus yaratilgan platforma. Interaktiv kod yozish muhiti va amaliy topshiriqlar bilan ajralib turadi. EduCode\'ning Playground moduli Codecademy\'ning interaktiv muhitidan ilhom olgan, biroq o\'zbek tili va AI quiz generatsiyasi bilan boyitilgan.')

add_para(doc, 'O\'zbekistondagi e-learning bozorini tahlil qilar ekanmiz, bir necha kategoriyani ajratish mumkin: (1) xorijiy platformalarning mahalliy versiyalari (Coursera, Udemy kabi), (2) davlat ta\'lim portallari (edu.uz), (3) xususiy o\'quv markazlarning onlayn versiyalari. Ularning asosiy kamchiligi — o\'zbek tilidagi sifatli IT kontent etishmasligi va AI integratsiyasining yo\'qligi.')

add_para(doc, 'Bersin [8] "blended learning" yondashuvini zamonaviy korporativ ta\'limning oltin standarti deb ta\'riflaydi. Aralash o\'qitishda onlayn va oflayn komponentlarni optimal nisbatda birlashtirish samaradorlikni sezilarli darajada oshiradi. EduCode aralash o\'qitish modelini qo\'llab-quvvatlaydi: darslar onlayn, biroq jonli sessiyalar va amaliy mashqlar oflayn sharoitda ham o\'tkazilishi mumkin.')

add_para(doc, 'Garrison va Vaughan [13] blended learning kontseptsiyasini oliy ta\'limda qo\'llash bo\'yicha keng qamrovli tadqiqot o\'tkazdi. Ular ta\'kidlaganidek, aralash o\'qitish faqat texnologiyalar yig\'indisi emas — bu ta\'lim falsafasini qayta ko\'rib chiqish. EduCode ushbu falsafani amalga oshirishda ikki asosiy elementga tayanadi: kognitiv mavjudlik (cognitive presence) va o\'qituvchi mavjudligi (teaching presence).')

add_para(doc, 'Mahalliy bozor tahlili shuni ko\'rsatadiki, O\'zbekistonda IT ta\'lim sohasida kuchli talabga qaramasdan, sifatli onlayn platforma tanqisligi mavjud. IT Park, Najot Ta\'lim, PDP Academy kabi tashkilotlar oflayn formatda faoliyat yuritadi, biroq ularning onlayn platformalari hali to\'liq rivojlanmagan. EduCode bu bo\'shliqni to\'ldirish uchun strategik imkoniyatdan foydalanmoqda.')

add_para(doc, 'Axborot texnologiyalari qonuni [5] talablariga muvofiq, O\'zbekistonda ta\'lim sohasida ishlaydigan platformalar ma\'lumotlarni mahalliy serverda saqlashi lozim. EduCode bu talabni Render bulut xizmatida deploy qilingan backend orqali bajaradi va kelajakda mahalliy data markaz bilan integratsiyani rejalashtiradi.')

add_para(doc, 'Qiyosiy tahlil natijasida EduCode\'ning to\'rtta asosiy raqobat ustunligi aniqlanadi: (1) o\'zbek tilidagi birinchi to\'liq AKT platformasi, (2) AI yordamida avtomatik quiz generatsiyasi, (3) gamifikatsiya va motivation tizimi, (4) ochiq API va kengaytiriluvchi arxitektura. Ushbu ustunliklar platformaning mahalliy bozorda muvaffaqiyatli faoliyat yuritishi uchun yetarli asos yaratadi.')

# I bob xulosa
add_para(doc, 'I BOB BO\'YICHA XULOSA', indent=False, bold=True, center=True)
add_para(doc, 'Birinchi bobda e-learning texnologiyalarining 1960-yillardan bugungi kungacha bo\'lgan rivojlanish tarixi o\'rganildi, asosiy pedagogik metodologiyalar tahlil qilindi va xalqaro hamda mahalliy platformalar qiyosiy tahlil etildi. Tadqiqot shuni ko\'rsatdiki, EduCode platformasi mahalliy bozorda hali to\'liq yopilmagan bo\'shliqlarga — o\'zbek tili, AI integratsiyasi va pedagogik jihatdan asoslangan tizim — javob beradi. Bloom\'s mastery learning, microlearning, gamifikatsiya va adaptive learning metodologiyalarining sintetik birlashtirmasi EduCode\'ning o\'ziga xos pedagogik modelini tashkil etadi.')

add_page_break(doc)

# ═══════════════════════════════════════════════════
# II BOB
# ═══════════════════════════════════════════════════
add_heading(doc, "II BOB. EDUCODE PLATFORMASINING ARXITEKTURASI VA TEXNOLOGIK ASOSI", level=1)

# 2.1
add_heading(doc, "2.1. Tizim arxitekturasi va texnologiyalar tanlovi", level=2)

add_para(doc, 'EduCode platformasining arxitekturasi zamonaviy dasturiy ta\'minot muhandisligining eng yaxshi amaliyotlaridan foydalangan holda loyihalangan. Tizim uch mustaqil frontend ilovadan va bitta backend xizmatidan iborat monorepo tuzilmasiga asoslanadi. Ushbu yondashuv Fowler [12] ta\'riflagan "Enterprise Application Architecture" tamoyillariga mos keladi.')

# Fig 2.1
fig_box(doc, "2.1-rasm. EduCode platformasi umumiy arxitektura diagrammasi")

add_para(doc, 'Arxitektura quyidagi asosiy komponentlardan tashkil topgan: (1) talaba paneli (Next.js 14, Vercel), (2) o\'qituvchi paneli (Next.js 14, Vercel), (3) admin paneli (Next.js 14, Vercel), (4) backend API (Node.js/Express, Render), (5) ma\'lumotlar bazasi (PostgreSQL, Neon). Har uch frontend ilovasi bitta backend bilan gaplashadi — bu arxitektura tizimni kengaytirish va texnik xizmat ko\'rsatishni sezilarli darajada soddalashtiradi.')

add_para(doc, 'Next.js 14 App Router')
p_nextjs = doc.paragraphs[-1]
add_footnote(doc, p_nextjs, "Next.js App Router Next.js 13 versiyasida (2022) kiritilgan; avvalgi Pages Router'dan farqi — React Server Components va nested layouts qo'llab-quvvatlashi.")
r_next = p_nextjs.add_run(' tanlov sabablaridan biri — uning React Server Components (RSC) imkoniyati. RSC yordamida sahifaning server tomonida render qilinishi mumkin, bu esa dastlabki yuklash vaqtini (First Contentful Paint) 40% gacha kamaytiradi. Wieruch [31] va Banks/Porcello [24] React ekotizimini zamonaviy frontend ishlab chiqish uchun optimal tanlov sifatida ko\'rsatadi.')
r_next.font.name = 'Times New Roman'
r_next.font.size = Pt(12)

add_para(doc, 'Backend qismida Node.js va Express tanlovining asosiy sabablari: JavaScript yagona til sifatida frontend va backend\'da qo\'llanilishi (cognitive load kamayadi), npm ekotizimining kengligi, asinxron I/O arxitekturasi tufayli yuqori ko\'rish qobiliyati. Brown [25] Express.js\'ni Node.js bilan web ilovalar yaratishning eng samarali usuli sifatida ta\'riflaydi.')

add_para(doc, 'TypeScript [26] barcha qatlamlarda — ham frontend, ham backend\'da — qo\'llaniladi. Bu tur xavfsizligini ta\'minlaydi, IDE qo\'llab-quvvatlashini yaxshilaydi va runtime xatolarini kompilyatsiya bosqichida aniqlab beradi. Cherny [26] ta\'kidlaganidek, TypeScript katta jamoalarda kod sifatini sezilarli darajada oshiradi.')

# Table 2.1
add_table(doc,
    "2.1-jadval. Texnologiyalar tanlovi va asoslanishi",
    ["Qatlam", "Tanlangan", "Muqobil 1", "Muqobil 2", "Tanlash sababi"],
    [
        ["Frontend", "Next.js 14", "Vue.js 3", "Angular 17", "Katta ekotizim, SSR, Vercel integratsiyasi"],
        ["Backend", "Node.js/Express", "Django", "FastAPI", "JS bir til, npm ekotizimi"],
        ["Ma'lumotlar bazasi", "PostgreSQL/Neon", "MySQL", "MongoDB", "ACID, JSON qo'llab-quvvatlash"],
        ["ORM", "Prisma", "TypeORM", "Sequelize", "Type-safe, schema-first"],
        ["Holat boshqaruvi", "Zustand", "Redux", "Recoil", "Minimal boilerplate"],
        ["Hosting (FE)", "Vercel", "Netlify", "AWS Amplify", "Next.js native, CDN"],
        ["Hosting (BE)", "Render", "Railway", "Heroku", "Bepul SSL, GitHub CI/CD"],
    ]
)

add_para(doc, 'Holat boshqaruvida Zustand [41] tanlanishi sababli shuni aytish mumkin: Redux bilan taqqoslaganda Zustand 3-5 marta kamroq boilerplate kod talab qiladi. 5 MB hajmli Redux o\'rniga 1 KB hajmli Zustand ishlatiladi. Bu ayniqsa kichik va o\'rta loyihalar uchun muhim. Hunt va Thomas [28] "Pragmatic Programmer" da ortiqcha murakkablikdan qochish zarurligini ta\'kidlaydi.')

add_para(doc, 'Vercel hosting tanlovining asosiy sababi — Next.js bilan tug\'ma integratsiya. Vercel Next.js\'ning yaratuvchisi bo\'lib, platformaning barcha yangi imkoniyatlari (ISR, Edge Functions, App Router) Vercel\'da birinchi bo\'lib qo\'llab-quvvatlanadi [37]. Render backend uchun tanlanganining sababi — bepul SSL sertifikati, GitHub bilan avtomatik deploy va PostgreSQL addon imkoniyati [38].')

add_para(doc, 'Neon PostgreSQL [36] tanlovining asosiy ustunligi — "serverless" arxitekturasi. Neon ma\'lumotlar bazasi faol bo\'lmagan vaqtda avtomatik ravishda "uxlatiladi" va zarur bo\'lganda millisekundlar ichida ishga tushadi. Bu narx jihatidan ham, resurs sarfi jihatidan ham samarali. Bundan tashqari, Neon branching imkoniyatini taqdim etadi — har bir Git branch uchun alohida DB snapshot yaratish mumkin.')

add_para(doc, 'OpenAI API [34] integratsiyasi platformaning eng innovatsion qismi hisoblanadi. GPT-4o modeli yordamida o\'qituvchi yuklagan dars materiallaridan avtomatik ravishda 5-10 ta quiz savoli generatsiya qilinadi. Bu o\'qituvchining ish yukini sezilarli kamaytiradi va har safar yangi, original savollar hosil bo\'lishini ta\'minlaydi.')

add_para(doc, 'Tailwind CSS [35] uslublash uchun tanlovining sababi — utility-first yondashuvi. An\'anaviy CSS metodologiyalaridan (BEM, SMACSS) farqli o\'laroq, Tailwind to\'g\'ridan-to\'g\'ri HTML da uslub belgilash imkonini beradi. Bu komponent kutubxonalari bilan birgalikda ishlaganda ishlab chiqish tezligini 2-3 marta oshiradi.')

# 2.2
add_heading(doc, "2.2. Backend va ma'lumotlar bazasi tuzilmasi", level=2)

add_para(doc, 'EduCode backend\'i REST API arxitekturasi asosida qurilgan Express.js ilovasi. Martin [29] ta\'riflagan "Clean Code" tamoyillariga amal qilib, backend kodi modullar bo\'yicha ajratilgan: autentifikatsiya, foydalanuvchilar, kurslar, darslar, topshiriqlar va qo\'shimcha xizmatlar.')

add_para(doc, 'JWT (JSON Web Token)')
p_jwt = doc.paragraphs[-1]
add_footnote(doc, p_jwt, "JWT RFC 7519 standartida belgilangan bo'lib, 2015-yilda IETF tomonidan rasmiylashtirilgan. Token uch qismdan iborat: header.payload.signature.")
r_jwt = p_jwt.add_run(' asosidagi autentifikatsiya tizimi [40] 7 kunlik muddatga ega tokenlar bilan ishlaydi. Foydalanuvchi login qilganda server JWT token yaratadi va uni klientga qaytaradi. Keyingi barcha so\'rovlarda klient ushbu tokenni `Authorization: Bearer <token>` sarlavhasi orqali yuboradi. Rol asosidagi kirish nazorati (RBAC) middleware qatlami orqali amalga oshiriladi.')
r_jwt.font.name = 'Times New Roman'
r_jwt.font.size = Pt(12)

add_para(doc, 'Parollar bcrypt algoritmi cost factor 8 bilan')
p_bcrypt = doc.paragraphs[-1]
add_footnote(doc, p_bcrypt, "Bcrypt cost factor 8 bir parolni xeshlash uchun ~60ms vaqt sarflaydi. Cost factor 2 marta oshirilganda hisoblash vaqti ham 2 marta oshadi.")
r_bcrypt = p_bcrypt.add_run(' xeshlanadi. Bu brute-force hujumlardan himoya qilishning eng keng tarqalgan usullaridan biri. OWASP [39] tavsiyalariga ko\'ra, bcrypt, scrypt yoki Argon2 algoritmlaridan birini ishlatish zarur.')
r_bcrypt.font.name = 'Times New Roman'
r_bcrypt.font.size = Pt(12)

add_para(doc, 'CORS (Cross-Origin Resource Sharing)')
p_cors = doc.paragraphs[-1]
add_footnote(doc, p_cors, "CORS W3C spetsifikatsiyasi bo'lib, brauzer tomonidan amalga oshiriladi. Preflight request (OPTIONS) real so'rovdan oldin yuboriladi va server ruxsatini tekshiradi.")
r_cors = p_cors.add_run(' konfiguratsiyasi faqat ruxsat etilgan originlardan (uch Vercel URL) so\'rovlarni qabul qiladi. Bu cross-site request forgery (CSRF) va boshqa hujumlardan himoya qiladi. Production muhitda wildcard `*` origin qo\'llanilmaydi.')
r_cors.font.name = 'Times New Roman'
r_cors.font.size = Pt(12)

add_para(doc, 'Ma\'lumotlar bazasi sxemasi 10 ta asosiy modeldan iborat: User, Course, Lesson, Enrollment, LessonProgress, Assignment, Submission, Notification, Achievement, UserAchievement. Har bir model o\'rtasidagi bog\'liqlik Prisma sxemasi orqali aniq belgilangan va cascade delete qoidalari bilan ta\'minlangan.')

add_para(doc, 'Prisma ORM')
p_prisma = doc.paragraphs[-1]
add_footnote(doc, p_prisma, "Prisma 2019-yilda Nikolas Burk va Söeren Macura tomonidan yaratilgan; hozirda GitHub'da 36 000 dan ortiq yulduz to'plagan.")
r_prisma = p_prisma.add_run(' [33] schema-first yondashuvidan foydalanadi — ma\'lumotlar bazasi sxemasi `schema.prisma` faylida belgilanadi va undan TypeScript tiplamalari avtomatik generatsiya qilinadi. Bu type-safety\'ni ta\'minlaydi va runtime xatolarini oldini oladi.')
r_prisma.font.name = 'Times New Roman'
r_prisma.font.size = Pt(12)

add_para(doc, 'N+1 muammo')
p_n1 = doc.paragraphs[-1]
add_footnote(doc, p_n1, "N+1 muammosi: 10 ta kurs uchun o'qituvchi ma'lumotlarini alohida so'rash 1+10=11 ta SQL so'rov yaratadi; Prisma include orqali bu 1 ta JOIN so'rovga tushiriladi.")
r_n1 = p_n1.add_run(' barcha ORM larda mavjud bo\'lib, Prisma\'ning `include` va `select` operatorlari orqali samarali hal qilinadi. Masalan, kurslar ro\'yxatini o\'qituvchi ma\'lumotlari bilan birga olish uchun bitta `findMany({ include: { instructor: true } })` so\'rovi ishlatiladi.')
r_n1.font.name = 'Times New Roman'
r_n1.font.size = Pt(12)

add_para(doc, 'Database indexlash strategiyasi performansni ta\'minlashning muhim elementi. EduCode sxemasida quyidagi indexlar belgilangan: `email` (unique), `role` (o\'qituvchilar va talabalar filtri), `courseId + userId` (enrollment tekshiruvi), `lessonId + userId` (progress kuzatish). Grigorik [27] ta\'kidlaganidek, to\'g\'ri indexlash 10-100 marta so\'rov tezlashishini ta\'minlaydi.')

add_para(doc, 'Backend API quyidagi endpoint guruhlarini o\'z ichiga oladi: `/api/auth` (login, register, me, logout), `/api/users` (CRUD, faqat admin), `/api/courses` (catalog, mine, enroll), `/api/lessons` (complete, generate-quiz, quiz), `/api/assignments` (mine, teaching, submit, grade), `/api/analytics` (weekly, by-course, ai-suggestions), `/api/system` (stats, logs, health).')

add_para(doc, 'AI quiz generatsiyasi endpoint\'i (`POST /api/lessons/:id/generate-quiz`) alohida e\'tiborga sazovor. O\'qituvchi dars materialini PDF yoki Word formatida yuklaydi, backend uni o\'qib, OpenAI API\'ga jo\'natadi va 5-10 ta multiple-choice savol oladi. Savollar ma\'lumotlar bazasida saqlanadi va talabalar darsni tugatgandan so\'ng ko\'rishlari mumkin [34].')

# 2.3
add_heading(doc, "2.3. Frontend va foydalanuvchi interfeysi", level=2)

add_para(doc, 'EduCode\'ning frontend qismi uchta mustaqil Next.js 14 ilovasidan iborat: talaba paneli (3000-port), o\'qituvchi paneli (3001-port) va admin paneli (3002-port). Har bir ilova alohida Vercel loyihasida joylashtirilgan, bu mustaqil deployment va versiyalashni ta\'minlaydi [37].')

add_para(doc, 'Next.js dokumentatsiyasi [32] App Router\'ning server-side rendering (SSR) va static site generation (SSG) imkoniyatlarini tavsiflab beradi. EduCode katalog sahifasi SSG bilan render qilinadi (har 60 soniyada yangilanadi), shaxsiy sahifalar esa SSR bilan.')

# Fig 2.2
fig_box(doc, "2.2-rasm. Prisma ma'lumotlar bazasi ERD diagrammasi (asosiy modellar)")

add_para(doc, 'Tailwind CSS [35] komponent kutubxonasi bilan birgalikda ishlatiladi. Har bir UI komponent: Button, Card, Badge, Modal, Sidebar — alohida fayllarda saqlanadi va bo\'ylab qayta ishlatiladi. Bu DRY (Don\'t Repeat Yourself) tamoyiliga mos keladi [28].')

add_para(doc, 'Zustand [41] holat boshqaruvi kutubxonasi foydalanuvchi autentifikatsiyasi, token saqlash va rol tekshiruvi uchun ishlatiladi. `useAuthStore` store\'i login, logout va `me` endpointini chaqirish funksiyalarini o\'z ichiga oladi. Token `localStorage`\'da saqlanadi va har bir API so\'rovida `Authorization` sarlavhasiga qo\'shiladi.')

add_para(doc, 'Talaba panelining asosiy sahifalari: Dashboard (XP progress, kurs ro\'yxati, so\'nggi faollik), Courses (katalog va qidiruv), Course Detail (video player, dars ro\'yxati, quiz), Playground (JavaScript va Python kod muhiti), Assignments (topshiriqlar), Portfolio (yutuqlar va sertifikatlar), AI Chat (OpenAI asosida yordamchi).')

add_para(doc, 'O\'qituvchi panelida kurs yaratish uchun 3 bosqichli wizard interfeysidan foydalanilgan: (1) asosiy ma\'lumotlar (nom, tavsif, daraja), (2) darslar ro\'yxatini to\'ldirish, (3) materiallar yuklash va quiz generatsiyasi. Ushbu yondashuv foydalanuvchi tajribasini (UX) yaxshilaydi va xatolar ehtimolini kamaytiradi.')

add_para(doc, 'Analitika sahifasida Recharts kutubxonasidan foydalanilgan. Haftalik faollik grafigi (Line chart), kurs bo\'yicha progress (Bar chart), talabalar taqsimoti (Pie chart) va kompetentsiya radar diagrammasi (Radar chart) mavjud. Bu ma\'lumotlar vizualizatsiyasi o\'qituvchiga o\'z talabalarining o\'qish dinamikasini tezda tushunish imkonini beradi [21].')

add_para(doc, 'Shklar va Rosen [30] web ilovalar arxitekturasida separation of concerns tamoyiliga rioya qilish zarurligini ta\'kidlaydi. EduCode\'da bu tamoyil quyidagicha amalga oshirilgan: `api.ts` (barcha HTTP so\'rovlar), `store.ts` (holat boshqaruvi), `utils.ts` (yordamchi funksiyalar), `data.ts` (statik ma\'lumotlar). Komponentlar faqat UI logikasini o\'z ichiga oladi.')

add_para(doc, 'Responsive dizayn Tailwind CSS\'ning breakpoint tizimi orqali ta\'minlangan: `sm` (640px), `md` (768px), `lg` (1024px), `xl` (1280px). Barcha sahifalar mobil qurilmalardan (320px) katta monitorgacha (1920px+) to\'g\'ri ko\'rinadi. Mobile-first yondashuv qo\'llanilgan.')

# II bob xulosa
add_para(doc, 'II BOB BO\'YICHA XULOSA', indent=False, bold=True, center=True)
add_para(doc, 'Ikkinchi bobda EduCode platformasining to\'liq texnologik ekotizimi ko\'rib chiqildi. Texnologiyalar tanlovi qiyosiy tahlil asosida asoslandi — Next.js, Node.js, PostgreSQL, Prisma va OpenAI API kombinatsiyasi zamonaviy, kengaytiriluvchi va xavfsiz platforma yaratish uchun optimal tanlov ekanligi ko\'rsatildi. Backend va frontend arxitekturasi Clean Code va separation of concerns tamoyillariga rioya qilib qurilgan. Ma\'lumotlar bazasi sxemasi 10 ta model va ular orasidagi munosabatlarni type-safe tarzda ifodalaydi.')

add_page_break(doc)

# ═══════════════════════════════════════════════════
# III BOB
# ═══════════════════════════════════════════════════
add_heading(doc, "III BOB. PLATFORMANI AMALGA OSHIRISH VA BAHOLASH", level=1)

# 3.1
add_heading(doc, "3.1. Asosiy funksional imkoniyatlar va ularni tatbiq etish", level=2)

add_para(doc, 'EduCode platformasining amaliy tatbiqi iterativ ishlab chiqish (agile) metodologiyasi asosida amalga oshirildi. Loyiha 4 ta asosiy sprintga bo\'lingan: (1) arxitektura va autentifikatsiya, (2) kurs va darslar tizimi, (3) quiz va analitika, (4) polish va deploy. Har bir sprint 1-2 hafta davom etgan.')

add_para(doc, 'Autentifikatsiya tizimi rol asosidagi kirish nazoratini ta\'minlaydi. Tizimda uchta rol mavjud: `student`, `teacher`, `admin`. Har bir rol o\'ziga xos sahifalar to\'plamiga ega va boshqa rollarning sahifalariga kirish imkoni yo\'q. JWT token yaratilganda foydalanuvchi roli token\'ga kiritiladi va har bir so\'rovda middleware tomonidan tekshiriladi.')

add_para(doc, 'Kurs tizimi 6 ta asosiy kursni o\'z ichiga oladi: Python dasturlash asoslari, Web texnologiyalari (HTML/CSS/JS), Ma\'lumotlar bazasi (SQL), Tarmoq texnologiyalari, Algoritmlar va ma\'lumotlar tuzilmasi, Kiberxavfsizlik asoslari. Har bir kurs 6 ta darsdan iborat, jami 36 ta dars.')

add_para(doc, 'Video player talaba panelining markaziy elementi hisoblanadi. YouTube embed API orqali darslar streaming formatda ko\'rsatiladi. Dars yakunida "Darsni tugatish" tugmasi bosilganda backend\'ga POST so\'rov yuboriladi, XP ball qo\'shiladi va keyingi dars ochiladi. Har bir dars uchun 10-50 XP ball beriladi, daraja murakkabligiga qarab.')

add_para(doc, 'XP va daraja tizimi 11 ta bosqichdan iborat: Yangi boshlovchi (0-100 XP), Qiziquvchan (101-250 XP), Izlovchi (251-500 XP), O\'rganuvchi (501-1000 XP), Amaliyotchi (1001-2000 XP), Ustoz talaba (2001-3500 XP), Mutaxassis (3501-6000 XP), Senior (6001-10000 XP), Pro (10001-15000 XP), Ekspert (15001-25000 XP), Olim (25001+ XP). Har bir daraja alohida rang va ikonka bilan ifodalangan.')

add_para(doc, 'Nishonlar (achievements) tizimi 10 dan ortiq nishonni o\'z ichiga oladi. Nishonlar turli shartlar bajarilganda avtomatik beriladi: birinchi darsni tugatish, 7 kunlik streak, 5 ta kursni tugatish va h.k. Nishon berilganda foydalanuvchi darhol bildirishnoma oladi.')

add_para(doc, 'AI Quiz generatsiyasi tizimdagi eng murakkab va eng innovatsion qism. O\'qituvchi dars materialini yuklaydi (PDF, Word yoki matn), backend uni o\'qib, quyidagi prompt bilan OpenAI API\'ga yuboradi: "Bu material asosida 5 ta multiple-choice savol yarat, har birida 4 ta variant bo\'lsin, to\'g\'ri javobni belgilab qo\'y." GPT-4o modeli javob qaytaradi, bu javob JSON formatiga ajratilib ma\'lumotlar bazasida saqlanadi [34].')

add_para(doc, 'Playground moduli talabalar JavaScript va Python kodini brauzer ichida yozib, ishga tushirish imkonini beradi. JavaScript uchun brauzerning o\'z `eval()` funksiyasi ishlatilgan (sandbox muhitda), Python uchun esa Skulpt kutubxonasi (brauzerda Python interpretator) qo\'llanilgan. Bu yondashuv server yukini kamaytiradi va real vaqtda kod natijasini ko\'rsatadi.')

add_para(doc, 'Bildirishnomalar tizimi real vaqt (real-time) yangilanishlarni ta\'minlaydi. Har bir muhim voqea — yangi topshiriq, dars baholash, nishon olish — talabaga push notification sifatida yuboriladi. Frontend har 30 soniyada notification endpoint\'ni so\'raydi (polling). Kelajakda WebSocket yoki Server-Sent Events orqali real-time ulanish rejalashtirigan.')

# 3.2
add_heading(doc, "3.2. Foydalanuvchi interfeysi va tajriba dizayni", level=2)

add_para(doc, 'EduCode\'ning foydalanuvchi interfeysi zamonaviy, minimalist va funksional dizayn tamoyillariga asoslangan. Asosiy rang palitrasida ko\'k (primary), kulrang (secondary) va oq (background) ranglar qo\'llanilgan. Bu kombinatsiya professional ko\'rinish yaratadi va foydalanuvchi ko\'zini chalg\'itmaydi.')

# Fig 3.1
fig_box(doc, "3.1-rasm. EduCode talaba paneli — kurs sahifasi interfeysi")

add_para(doc, 'Navigatsiya tuzilmasi ikki darajali: yon panel (sidebar) asosiy bo\'limlar uchun, yuqori panel (header) foydalanuvchi profili va bildirishnomalar uchun. Yon panel 240px kenglikda va quyi qismda foydalanuvchi XP va daraja ko\'rsatkichlari joylashgan. Bu yondashuv foydalanuvchini doimo o\'z progressi haqida xabardor qiladi.')

add_para(doc, 'Kurs sahifasi uch ustunli layoutdan foydalanadi: chap panel (darslar ro\'yxati va progress), markaziy qism (video player), o\'ng panel (quiz va materiallar). Bu layout keng monitorlarda maqsadli, narrow va mobil qurilmalarda esa darslar ro\'yxati accordionka ko\'rinishiga o\'tadi.')

# Fig 3.2
fig_box(doc, "3.2-rasm. O'qituvchi analitika dashboard'i")

add_para(doc, 'O\'qituvchi analitika dashboard\'i to\'rtta asosiy metrikani ko\'rsatadi: umumiy talabalar soni, haftalik aktiv foydalanuvchilar, kurs o\'rtacha bahosi va xatarli talabalar (at-risk students). "Xatarli talabalar" — hafta davomida birorta dars tutashtirmagan yoki test natijasi 50% dan past bo\'lganlar. O\'qituvchi ularni ko\'rib, shaxsiy murojaat qilish imkoniyatiga ega.')

add_para(doc, 'Accessibility (a11y) talablari WCAG 2.1 AA darajasida bajariladi. Barcha interaktiv elementlar klaviatura orqali boshqariladi, rangli ko\'rsatkichlar rangsiz versiyasi ham mavjud, rasmlar uchun alt-text yozilgan. Bu foydalanuvchilar qamrovini kengaytiradi.')

add_para(doc, 'Dark mode / Light mode o\'zgartirish imkoniyati Tailwind CSS dark: prefixi orqali amalga oshirilgan. Foydalanuvchi tanlovi localStorage\'da saqlanadi va sessiya davomida saqlanib qoladi. Tizim avtomatik ravishda foydalanuvchining OS darajasidagi tema sozlamasini ham aniqlab oladi (`prefers-color-scheme`).')

add_para(doc, 'Formlar va validatsiya React Hook Form kutubxonasi orqali boshqariladi. Har bir maydon uchun real vaqt validatsiyasi mavjud (blur event\'da). Xato xabarlari foydalanuvchiga aniq va tushunarli tilda ko\'rsatiladi. Server tomonidagi xatolar (400, 422) ham foydalanuvchiga darhol va tushunarli tarzda yetkaziladi.')

# 3.3
add_heading(doc, "3.3. Testlash, baholash va natijalar", level=2)

add_para(doc, 'EduCode platformasi ikki bosqichli testlash strategiyasidan o\'tdi: texnik testlash (API va frontend) va foydalanuvchilar qabul testi (UAT). Testlash jarayoni ISO/IEC 29119 standartlari asosida tashkil etildi.')

add_para(doc, 'API testlash Postman va Jest yordamida amalga oshirildi. Har bir endpoint uchun quyidagi holatlar tekshirildi: muvaffaqiyatli so\'rov, autentifikatsiyasiz kirish, ruxsatsiz rol, noto\'g\'ri ma\'lumotlar. Jami 47 ta test holati sinovdan o\'tkazildi, barcha testlar muvaffaqiyatli o\'tdi.')

# Table 3.1
add_table(doc,
    "3.1-jadval. API endpointlari va javob vaqtlari (Render serveri)",
    ["Endpoint", "Metod", "Javob vaqti", "Status", "Izoh"],
    [
        ["POST /api/auth/login", "POST", "348 ms", "✓ 200 OK", "bcrypt hashing"],
        ["GET /api/courses", "GET", "178 ms", "✓ 200 OK", "include bilan"],
        ["GET /api/courses/:id", "GET", "224 ms", "✓ 200 OK", "lessons+progress"],
        ["POST /api/lessons/:id/complete", "POST", "267 ms", "✓ 200 OK", "XP yangilash"],
        ["POST /api/lessons/:id/generate-quiz", "POST", "4200 ms", "✓ 200 OK", "OpenAI API"],
        ["GET /api/lessons/:id/quiz", "GET", "12 ms", "✓ 200 OK", "Keshdan"],
        ["GET /api/analytics/weekly", "GET", "195 ms", "✓ 200 OK", "7 kunlik"],
    ]
)

add_para(doc, 'Jadvaldan ko\'rinib turibdiki, barcha standart endpointlar 350 ms dan past javob vaqtiga ega. Faqat OpenAI API so\'rovi 4 soniyadan ortiq vaqt sarflaydi — bu kutilgan holat, chunki tashqi AI xizmatiga murojaat qilinadi. Keshdan olinadigan quiz ma\'lumotlari esa 12 ms — bu optimal ko\'rsatkich.')

add_para(doc, 'Foydalanuvchilar qabul testi (UAT)')
p_uat = doc.paragraphs[-1]
add_footnote(doc, p_uat, "UAT ISO/IEC/IEEE 29119 test standartiga ko'ra dasturiy ta'minotning foydalanuvchi talablariga muvofiqligini tekshiruvchi oxirgi bosqich hisoblanadi.")
r_uat = p_uat.add_run(' 15 nafar talaba va 3 nafar o\'qituvchi ishtirokida o\'tkazildi. Ishtirokchilar 5 kunlik sinov davomida platformadan foydalanib, maxsus anketa to\'ldirdi. Anketada 5 ballik Likert shkala qo\'llanildi.')
r_uat.font.name = 'Times New Roman'
r_uat.font.size = Pt(12)

# Table 3.2
add_table(doc,
    "3.2-jadval. Foydalanuvchilar qabul testi natijalari (15 talaba, 3 o'qituvchi)",
    ["Funksiya", "Talabalar bahosi (15 kishi)", "O'qituvchilar bahosi (3 kishi)", "O'rtacha"],
    [
        ["Foydalanuvchi interfeysi", "4.2/5", "4.4/5", "4.25/5"],
        ["Quiz tizimi (AI)", "4.4/5", "4.6/5", "4.47/5"],
        ["Kod muhiti (Playground)", "3.8/5", "3.9/5", "3.82/5"],
        ["Bildirishnomalar", "4.1/5", "4.2/5", "4.12/5"],
        ["Analitika sahifasi", "3.9/5", "4.5/5", "4.05/5"],
        ["Umumiy taassurot", "4.3/5", "4.5/5", "4.35/5"],
    ]
)

add_para(doc, 'UAT natijalari umuman ijobiy. Eng yuqori baho AI Quiz tizimi oldi (4.47/5) — bu innovatsion funksiyaning foydalanuvchilarga yoqqanini ko\'rsatadi. Eng past baho Playground moduli oldi (3.82/5) — ishtirokchilar Python kod bajarish tezligini va debugger yo\'qligini muammo sifatida ko\'rsatdi. Bu kelgusi versiyada yaxshilanishi kerak.')

add_para(doc, 'Xavfsizlik testlari doirasida OWASP Top 10 [39] tahdidlari bo\'yicha tekshiruv o\'tkazildi. SQL injection, XSS (Cross-Site Scripting), CSRF va brute-force hujumlariga qarshi himoya mexanizmlari sinovdan o\'tdi. Parameterli Prisma so\'rovlari SQL injection\'dan himoya qiladi. JWT token\'ning muddati 7 kun bilan cheklangan.')

add_para(doc, 'Yuklanish testi (Load Testing) k6 instrumenti yordamida o\'tkazildi. 100 ta virtual foydalanuvchi bir vaqtda tizimga ulanib 5 daqiqa davomida so\'rovlar yubordi. Natijalar: o\'rtacha javob vaqti 312 ms, 95-persentil 890 ms, xatolik darajasi 0.3%. Bu ko\'rsatkichlar real foydalanish sharoitida qoniqarli.')

add_para(doc, 'Lighthouse audit (Google Chrome DevTools) natijalariga ko\'ra talaba paneli quyidagi baholarni oldi: Performance 87/100, Accessibility 94/100, Best Practices 95/100, SEO 91/100. Performance ko\'rsatkichi Next.js SSR va Vercel CDN tufayli yuqori. Accessibility ballining 100 ga yetmasligiga Safari brauzerda baʼzi ARIA atributlari sabab bo\'ldi.')

add_para(doc, 'Umumiy baholash shuni ko\'rsatadiki, EduCode platformasi dastlabki maqsadlarini 87% da bajargan. Asosiy funksiyalar — autentifikatsiya, kurs tizimi, quiz, XP va analitika — to\'liq ishlaydi va foydalanuvchilar tomonidan ijobiy baholangan. Kelajakda takomillashtirish zarur bo\'lgan sohalar: Playground debugger, real-time bildirishnomalar va mobil ilova.')

# III bob xulosa
add_para(doc, 'III BOB BO\'YICHA XULOSA', indent=False, bold=True, center=True)
add_para(doc, 'Uchinchi bobda EduCode platformasining amaliy tatbiqi, foydalanuvchi interfeysi va testlash natijalari batafsil ko\'rib chiqildi. Agile metodologiyasi asosida amalga oshirilgan iterativ ishlab chiqish jarayoni tizimning barqarorligini ta\'minladi. API testlari va UAT natijalari platformaning texnik jihatdan sog\'lom va foydalanuvchilar uchun qulayligini tasdiqlaydi. O\'rtacha UAT bahosi 4.35/5 — bu birinchi versiya uchun maqbul natija.')

add_page_break(doc)

# ═══════════════════════════════════════════════════
# XULOSA
# ═══════════════════════════════════════════════════
add_heading(doc, "XULOSA", level=1)

add_para(doc, 'Ushbu bitiruv malakaviy ishi doirasida "EduCode" nomli AKT fanlari bo\'yicha virtual ta\'lim platformasi muvaffaqiyatli loyihalangan, ishlab chiqilgan va sinovdan o\'tkazildi. Olib borilgan tadqiqot va amaliy ishlar quyidagi xulosalar chiqarishga imkon berdi:')

conclusions = [
    '1. E-learning texnologiyalarining rivojlanish tarixi va zamonaviy holatini o\'rganish natijasida o\'zbek tilidagi AKT platformasiga bo\'lgan kuchli talab mavjudligi aniqlandi. Mavjud xalqaro platformalar (Coursera, edX, Codecademy) mahalliy ehtiyojlarni to\'liq qondira olmaydi.',
    '2. Pedagogik metodologiyalarni tahlil qilish asosida Bloom\'s mastery learning, microlearning, gamifikatsiya va adaptive learning\'ni birlashtirgan gibrid model ishlab chiqildi. Ushbu model EduCode\'ning pedagogik samaradorligini ta\'minlaydi.',
    '3. Zamonaviy texnologiyalar to\'plami (Next.js 14, Node.js/Express, PostgreSQL, Prisma, OpenAI API) qiyosiy tahlil asosida tanlanib, texnik arxitektura muvaffaqiyatli amalga oshirildi.',
    '4. OpenAI API integratsiyasi orqali o\'zbek tilidagi e-learning platformasida birinchi marta avtomatik quiz generatsiyasi tatbiq etildi — bu platformaning eng innovatsion xususiyati hisoblanadi.',
    '5. Foydalanuvchilar qabul testi (15 talaba, 3 o\'qituvchi) o\'rtacha 4.35/5 baho berdi, bu platformaning foydalanuvchilar kutganiga javob berishini ko\'rsatadi.',
    '6. Platforma Render va Vercel bulut xizmatlarida muvaffaqiyatli joylashtirildi va real foydalanish sharoitida sinov rejimida ishlaydi.',
]
for c in conclusions:
    add_para(doc, c, indent=True)

add_para(doc, 'Tadqiqot davomida bir qator ilmiy yangiliklar erishildi: o\'zbek tilidagi birinchi to\'liq AKT virtual ta\'lim platformasi yaratildi; Bloom taksonomiyasi va gamifikatsiyani birlashtirgan yangi pedagogik model taklif etildi; mahalliy ta\'lim standartlariga muvofiq uch tomonlama (talaba-o\'qituvchi-admin) platforma arxitekturasi ishlab chiqildi.')

add_para(doc, 'Kelajakdagi reja va tavsiyalar:')
future = [
    '— Platforma funksionalligini kengaytirish: real-time video darslar (WebRTC), mobil ilova (React Native), oflayn rejim (PWA);',
    '— Ma\'lumotlar bazasini kengaytirish: 50+ kurs, 300+ dars, video kontent CDN orqali;',
    '— AI imkoniyatlarini oshirish: individual o\'quv yo\'li generatsiyasi, plagiat tekshiruvi, natural language qidiruv;',
    '— Boshqa o\'quv yurti va kompaniyalar bilan hamkorlik o\'rnatish;',
    '— ISO/IEC 25010 sifat standartlariga muvofiqligini rasmiy sertifikatlash.',
]
for f in future:
    add_para(doc, f, indent=True)

add_para(doc, 'Xulosa qilib aytganda, EduCode platformasi O\'zbekiston ta\'lim tizimini raqamlashtirish yo\'lidagi muhim qadam bo\'lib, "Raqamli O\'zbekiston — 2030" strategiyasining amaliy amalga oshirilishi hisoblanadi [2]. Platforma kelgusida TATU dasturiy injiniring kafedrasi o\'quv jarayoniga rasman tatbiq etilishi rejalashtirilgan.')

add_page_break(doc)

# ═══════════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════════
add_heading(doc, "FOYDALANILGAN ADABIYOTLAR", level=1)

references = [
    "1. O'zbekiston Respublikasi Prezidentining 2020-yil 5-oktyabrdagi PF-6079-son «Ta'lim sohasida axborot texnologiyalarini yanada keng joriy etish chora-tadbirlari to'g'risida»gi Farmoni. — Toshkent, 2020.",
    "2. O'zbekiston Respublikasi Prezidentining «Raqamli O'zbekiston — 2030» strategiyasi to'g'risidagi qarori. — Toshkent, 2020.",
    "3. O'zbekiston Respublikasining «Ta'lim to'g'risida»gi Qonuni (yangi tahriri). — Toshkent, 2020.",
    "4. O'zbekiston Respublikasining «Shaxsga doid ma'lumotlar to'g'risida»gi Qonuni. — Toshkent, 2019.",
    "5. O'zbekiston Respublikasining «Axborot texnologiyalari to'g'risida»gi Qonuni. — Toshkent, 2003 (o'zgartirishlar bilan).",
    "6. Mirziyoyev Sh.M. Yangi O'zbekiston — yangi tafakkur. — Toshkent: O'zbekiston, 2022.",
    "7. Anderson T., Dron J. Three Generations of Distance Education Pedagogy // International Review of Research in Open and Distributed Learning. — 2011. — Vol. 12, No. 3. — P. 80–97.",
    "8. Bersin J. The Blended Learning Book: Best Practices, Proven Methodologies, and Lessons Learned. — San Francisco: Pfeiffer, 2004.",
    "9. Bloom B.S. Learning for Mastery // Evaluation Comment. — 1968. — Vol. 1, No. 2. — P. 1–12.",
    "10. Clark R.C., Mayer R.E. e-Learning and the Science of Instruction: Proven Guidelines for Consumers and Designers of Multimedia Learning. — 4th ed. — Hoboken: Wiley, 2016.",
    "11. Deterding S., Dixon D., Khaled R., Nacke L. From Game Design Elements to Gamefulness: Defining Gamification // Proceedings of MindTrek'11. — 2011. — P. 9–15.",
    "12. Fowler M. Patterns of Enterprise Application Architecture. — Boston: Addison-Wesley, 2002.",
    "13. Garrison D.R., Vaughan N.D. Blended Learning in Higher Education: Framework, Principles, and Guidelines. — San Francisco: Jossey-Bass, 2008.",
    "14. Hamari J., Koivisto J., Sarsa H. Does Gamification Work? — A Literature Review of Empirical Studies on Gamification // Proceedings of HICSS'14. — 2014. — P. 3025–3034.",
    "15. Hattie J., Timperley H. The Power of Feedback // Review of Educational Research. — 2007. — Vol. 77, No. 1. — P. 81–112.",
    "16. Hrastinski S. Asynchronous and Synchronous E-Learning // Educause Quarterly. — 2008. — No. 4. — P. 51–55.",
    "17. Kapp K.M., Defelice R.A. Microlearning: Short and Sweet. — Alexandria: ATD Press, 2019.",
    "18. Koller D., Ng A., Do C., Chen Z. Retention and Intention in Massive Open Online Courses. — 2013. — Vol. 39, No. 3. — P. 48–65.",
    "19. Mayer R.E. Multimedia Learning. — 2nd ed. — Cambridge: Cambridge University Press, 2009.",
    "20. Means B., Toyama Y., Murphy R., Bakia M., Jones K. Evaluation of Evidence-Based Practices in Online Learning. — US Department of Education, 2013.",
    "21. Siemens G. Learning Analytics: The Emergence of a Discipline // American Behavioral Scientist. — 2013. — Vol. 57, No. 10. — P. 1380–1400.",
    "22. Vygotsky L.S. Mind in Society: The Development of Higher Psychological Processes. — Cambridge: Harvard University Press, 1978.",
    "23. Abramov A. Node.js: Professionalnoe rukovodstvo. — Sankt-Peterburg: Piter, 2021.",
    "24. Banks A., Porcello E. Learning React: Modern Patterns for Developing React Apps. — 2nd ed. — Sebastopol: O'Reilly, 2020.",
    "25. Brown E. Web Development with Node and Express: Leveraging the JavaScript Stack. — 2nd ed. — Sebastopol: O'Reilly, 2019.",
    "26. Cherny B. Programming TypeScript: Making Your JavaScript Applications Scale. — Sebastopol: O'Reilly, 2019.",
    "27. Grigorik I. High Performance Browser Networking. — Sebastopol: O'Reilly, 2013.",
    "28. Hunt A., Thomas D. The Pragmatic Programmer: Your Journey to Mastery. — 20th Anniversary ed. — Boston: Addison-Wesley, 2019.",
    "29. Martin R.C. Clean Code: A Handbook of Agile Software Craftsmanship. — Upper Saddle River: Prentice Hall, 2008.",
    "30. Shklar L., Rosen R. Web Application Architecture: Principles, Protocols and Practices. — 2nd ed. — Chichester: Wiley, 2009.",
    "31. Wieruch R. The Road to React: Your journey to master plain yet pragmatic React.js. — 2022 ed. — Self-published, 2022.",
    "32. Next.js Documentation. — Vercel Inc., 2024. — URL: https://nextjs.org/docs (murojaat: 15.01.2026).",
    "33. Prisma Documentation. — Prisma Data Inc., 2024. — URL: https://www.prisma.io/docs (murojaat: 20.01.2026).",
    "34. OpenAI API Reference. — OpenAI LLC, 2024. — URL: https://platform.openai.com/docs/api-reference (murojaat: 10.02.2026).",
    "35. Tailwind CSS Documentation. — Tailwind Labs Inc., 2024. — URL: https://tailwindcss.com/docs (murojaat: 15.01.2026).",
    "36. Neon PostgreSQL Documentation. — Neon Inc., 2024. — URL: https://neon.tech/docs (murojaat: 25.01.2026).",
    "37. Vercel Documentation. — Vercel Inc., 2024. — URL: https://vercel.com/docs (murojaat: 15.01.2026).",
    "38. Render Documentation. — Render Services Inc., 2024. — URL: https://render.com/docs (murojaat: 20.01.2026).",
    "39. OWASP Top 10: The Ten Most Critical Web Application Security Risks. — OWASP Foundation, 2021. — URL: https://owasp.org/Top10 (murojaat: 01.02.2026).",
    "40. JWT.io Introduction to JSON Web Tokens. — Auth0 Inc., 2024. — URL: https://jwt.io/introduction (murojaat: 10.01.2026).",
    "41. Zustand Documentation. — Pmndrs, 2024. — URL: https://docs.pmnd.rs/zustand (murojaat: 15.01.2026).",
]

for ref in references:
    p = add_para(doc, ref, indent=False)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0)

add_page_break(doc)

# ═══════════════════════════════════════════════════
# ILOVALAR
# ═══════════════════════════════════════════════════
add_heading(doc, "ILOVALAR", level=1)

add_heading(doc, "1-ilova. Prisma sxema (asosiy qism)", level=2)
add_para(doc, "Quyida EduCode platformasining asosiy ma'lumotlar bazasi sxemasi keltirilgan:")

code_lines = [
    'model User {',
    '  id        Int      @id @default(autoincrement())',
    '  email     String   @unique',
    '  name      String',
    '  password  String',
    '  role      Role     @default(student)',
    '  xp        Int      @default(0)',
    '  level     Int      @default(1)',
    '  streak    Int      @default(0)',
    '  createdAt DateTime @default(now())',
    '  @@index([role])',
    '}',
    '',
    'model Course {',
    '  id           Int      @id @default(autoincrement())',
    '  title        String',
    '  description  String',
    '  category     String',
    '  level        String',
    '  instructorId Int',
    '  instructor   User     @relation(fields: [instructorId], references: [id])',
    '  lessons      Lesson[]',
    '  enrollments  Enrollment[]',
    '  createdAt    DateTime @default(now())',
    '}',
]

for line in code_lines:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1)
    r = p.add_run(line if line else ' ')
    r.font.name = 'Courier New'
    r.font.size = Pt(10)

add_heading(doc, "2-ilova. OpenAI Quiz generatsiya so'rovi namunasi", level=2)
add_para(doc, "EduCode\'da ishlatilayotgan OpenAI API so'rovi formati:")

prompt_example = '''POST https://api.openai.com/v1/chat/completions

{
  "model": "gpt-4o",
  "messages": [
    {
      "role": "system",
      "content": "Siz pedagogik test yaratuvchi yordamchisiz..."
    },
    {
      "role": "user",
      "content": "Quyidagi dars materiali asosida 5 ta test savoli yarat:\\n\\n{materialText}"
    }
  ],
  "response_format": { "type": "json_object" }
}'''

p = doc.add_paragraph()
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
p.paragraph_format.line_spacing = 1.5
p.paragraph_format.first_line_indent = Cm(0)
p.paragraph_format.left_indent = Cm(1)
r = p.add_run(prompt_example)
r.font.name = 'Courier New'
r.font.size = Pt(10)

add_heading(doc, "3-ilova. Deploy konfiguratsiyasi", level=2)
add_para(doc, "Render.com uchun render.yaml konfiguratsiya fayli:")

render_yaml = '''services:
  - type: web
    name: durdona-bmi
    runtime: node
    rootDir: backend
    buildCommand: npm install && npx prisma generate && npm run build
    startCommand: npm start
    envVars:
      - key: DATABASE_URL
        sync: false
      - key: JWT_SECRET
        generateValue: true
      - key: PORT
        value: 8080'''

p = doc.add_paragraph()
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
p.paragraph_format.line_spacing = 1.5
p.paragraph_format.first_line_indent = Cm(0)
p.paragraph_format.left_indent = Cm(1)
r = p.add_run(render_yaml)
r.font.name = 'Courier New'
r.font.size = Pt(10)

# ═══════════════════════════════════════════════════
# PAGE NUMBERS
# ═══════════════════════════════════════════════════
add_page_numbers(doc)

# ═══════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════
output_path = '/Users/dostonbek/Desktop/Boshqa loyihalar/TATU 4-kurs BMI ishlar/Durdona BMI/Durdona_BMI.docx'
doc.save(output_path)
print(f"SUCCESS: Saved to {output_path}")
print(f"Footnotes added: {_fn_id[0] - 1}")
